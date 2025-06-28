import os
import json
import re
import html as html_mod
from bs4 import BeautifulSoup
from flask import Flask, send_file, redirect, Response, render_template_string
import io
try:
    import requests
except ImportError:
    requests = None
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None

# Configuration: API endpoints
SERIES_API_URL = 'https://www.testranking.in/admin/api/get-tab-package-series-v1/615/0/1/0/-1/2650169/0/1000/561'
QUESTION_API_PATTERN = 'https://www.testranking.in/admin/api/questions-solutions-new/{series_id}/en'
INPUT_FILE = None

# Utility functions

def fetch_json_from_url(url):
    if not url:
        raise ValueError('URL is required')
    if not requests:
        raise RuntimeError('requests module is required')
    resp = requests.get(url, timeout=300)  # Increased timeout
    resp.raise_for_status()
    try:
        return resp.json()
    except Exception:
        return extract_json_from_html_wrapper(resp.text)


def fetch_raw_json(series_id=None):
    if INPUT_FILE:
        if not os.path.exists(INPUT_FILE):
            raise FileNotFoundError(f'Input file not found: {INPUT_FILE}')
        content = open(INPUT_FILE, encoding='utf-8').read()
        try:
            return json.loads(content)
        except Exception:
            return extract_json_from_html_wrapper(content)
    else:
        if not series_id:
            raise ValueError('series_id is required when not using INPUT_FILE')
        url = QUESTION_API_PATTERN.format(series_id=series_id)
        return fetch_json_from_url(url)


def extract_json_from_html_wrapper(html_content):
    m = re.search(r'<body[^>]*>(.*)</body>', html_content, flags=re.DOTALL | re.IGNORECASE)
    inner = m.group(1).strip() if m else html_content.strip()
    inner = html_mod.unescape(inner)
    try:
        return json.loads(inner)
    except Exception:
        fixed = inner.replace('\\"', '"')
        return json.loads(fixed)


def parse_html_field(s):
    if not s:
        return '', []
    soup = BeautifulSoup(s, 'html.parser')
    text = soup.get_text(separator=' ', strip=True)
    text = re.sub(r'\s+', ' ', text).strip()
    images = []
    for img in soup.find_all('img'):
        src = img.get('src')
        if src:
            images.append(src)
    return text, images


def process_data(raw_json, lang='en'):
    cleaned = []
    labels = ['a', 'b', 'c', 'd', 'e']
    for section_entry in raw_json.get('data', []):
        sec_id = section_entry.get('section_id') or section_entry.get('sec_id')
        sec_name = section_entry.get('section_name') or section_entry.get('sec_name') or str(sec_id)
        for q_list in section_entry.get('all_questions', {}).values():
            for q in q_list:
                qid = q.get('qid')
                topic_id = q.get('topic_id')
                question_field = f'question_{lang}'
                q_html = q.get(question_field) or ''
                question_text, question_images = parse_html_field(q_html)
                options = []
                for i in range(1, 6):
                    opt_field = f'option_{lang}_{i}'
                    opt_html = q.get(opt_field)
                    if opt_html:
                        opt_text, opt_images = parse_html_field(opt_html)
                        if opt_text or opt_images:
                            options.append({'text': opt_text, 'images': opt_images})
                ans_idx = None
                ans_text = ''
                ar = q.get('answer_en')
                if ar:
                    try:
                        i0 = int(ar)
                        if 1 <= i0 <= len(options):
                            ans_idx = i0 - 1
                            # Set ans_text to label only (a, b, c, d, e)
                            ans_text = labels[ans_idx] if ans_idx < len(labels) else str(i0)
                        else:
                            ahtml = q.get(f'answer_{lang}') or ar
                            ans_text, _ = parse_html_field(ahtml)
                    except Exception:
                        ahtml = q.get(f'answer_{lang}') or ar
                        ans_text, _ = parse_html_field(ahtml)
                sol_field = f'solution_{lang}'
                sol_html = q.get(sol_field) or ''
                solution_text, solution_images = parse_html_field(sol_html)
                entry = {
                    'qid': qid,
                    'section_id': sec_id,
                    'section_name': sec_name,
                    'topic_id': topic_id,
                    'question_text': question_text,
                    'question_images': question_images,
                    'options': options,
                    'answer_index': ans_idx,
                    'answer_text': ans_text,
                    'solution_text': solution_text,
                    'solution_images': solution_images
                }
                cleaned.append(entry)
    return cleaned


def group_by_section(cleaned):
    sections = {}
    for entry in cleaned:
        sid = entry.get('section_id') or ''
        sname = entry.get('section_name') or str(sid)
        if sid not in sections:
            sections[sid] = {'name': sname, 'questions': []}
        sections[sid]['questions'].append(entry)
    return sections


def create_docx(cleaned_entries, base_url=None):
    if not Document:
        raise RuntimeError('python-docx is required. Install with pip install python-docx')
    doc = Document()
    for entry in cleaned_entries:
        p = doc.add_paragraph()
        p.add_run(f"[Q] {entry['question_text']}")
        for img_url in entry.get('question_images', []):
            img_data = fetch_image_bytes(img_url, base_url)
            if img_data:
                doc.add_picture(io.BytesIO(img_data), width=Inches(4))
        labels = ['a', 'b', 'c', 'd', 'e']
        for idx, opt in enumerate(entry.get('options', [])):
            p_opt = doc.add_paragraph()
            label = labels[idx] if idx < len(labels) else str(idx+1)
            p_opt.add_run(f"({label}) {opt.get('text','')}")
            for img_url in opt.get('images', []):
                img_data = fetch_image_bytes(img_url, base_url)
                if img_data:
                    doc.add_picture(io.BytesIO(img_data), width=Inches(4))
        p_ans = doc.add_paragraph()
        p_ans.add_run(f"[ANS] {entry.get('answer_text','')}")
        p_sol = doc.add_paragraph()
        p_sol.add_run(f"[SOL] {entry.get('solution_text','')}")
        for img_url in entry.get('solution_images', []):
            img_data = fetch_image_bytes(img_url, base_url)
            if img_data:
                doc.add_picture(io.BytesIO(img_data), width=Inches(4))
        doc.add_page_break()
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def fetch_image_bytes(img_url, base_url=None):
    if not img_url:
        return None
    if img_url.startswith('//'):
        img_url = 'https:' + img_url
    elif img_url.startswith('/') and base_url:
        from urllib.parse import urljoin
        img_url = urljoin(base_url, img_url)
    if not requests:
        return None
    try:
        resp = requests.get(img_url, timeout=15)  # 120 ko 15-20 sec kar dein
        resp.raise_for_status()
        return resp.content
    except Exception:
        return None

# Initialize Flask
app = Flask(__name__)

# Base template with placeholder for content
BASE_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <title>{{ title }}</title>
    <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css">
</head>
<body>
<nav class="navbar navbar-expand-lg navbar-light bg-light">
  <a class="navbar-brand" href="/">Test Downloader</a>
</nav>
<div class="container my-4">
  {{ content|safe }}
</div>
</body>
</html>
'''

@app.route('/')
def index():
    if not requests:
        return 'requests module required', 500
    try:
        data = fetch_json_from_url(SERIES_API_URL)
    except Exception as e:
        return f'Error fetching series: {e}', 500
    series_list = data.get('data', [])
    index_content = '''
<h1 class="mb-4">Available Series</h1>
<div class="list-group">
{% for series in series_list %}
  <a href="/series/{{ series.series_id }}" class="list-group-item list-group-item-action">
    {{ series.series_name }} {% if series.is_attempted == '1' %}<span class="badge badge-success">Attempted</span>{% endif %}
  </a>
{% endfor %}
</div>
'''
    content = render_template_string(index_content, series_list=series_list)
    return render_template_string(BASE_TEMPLATE, title='Series List', content=content)

@app.route('/series/<series_id>')
def series_page(series_id):
    if not requests:
        return 'requests module required', 500
    try:
        data = fetch_json_from_url(SERIES_API_URL)
    except Exception as e:
        return f'Error: {e}', 500
    series_list = data.get('data', [])
    series = next((s for s in series_list if str(s.get('series_id')) == str(series_id)), None)
    if not series:
        return 'Series not found', 404
    try:
        raw_json = fetch_raw_json(series_id)
    except Exception as e:
        return f'Error fetching questions: {e}', 500
    cleaned_en = process_data(raw_json, 'en')
    sections = group_by_section(cleaned_en)
    series_content = '''
<h1 class="mb-3">Series: {{ series.series_name }}</h1>
<p><a href="/">&laquo; Back to series list</a></p>
<ul class="list-inline">
  <li class="list-inline-item"><a href="/raw/{{ series.series_id }}" class="btn btn-primary">View Raw JSON</a></li>
  <li class="list-inline-item"><a href="/download/full/{{ series.series_id }}/en" class="btn btn-success">Download Full English</a></li>
  <li class="list-inline-item"><a href="/download/full/{{ series.series_id }}/hi" class="btn btn-success">Download Full Hindi</a></li>
</ul>
<h2 class="mt-4">Sections</h2>
<div class="list-group">
{% for sid, sec in sections.items() %}
  <div class="list-group-item">
    <strong>{{ sec.name }}</strong>
    <div class="btn-group btn-group-sm float-right" role="group">
      <a href="/download/section/{{ series.series_id }}/{{ sid }}/en" class="btn btn-outline-secondary">EN</a>
      <a href="/download/section/{{ series.series_id }}/{{ sid }}/hi" class="btn btn-outline-secondary">HI</a>
      <a href="/view/section/{{ series.series_id }}/{{ sid }}" class="btn btn-outline-primary">View Qs</a>
    </div>
  </div>
{% endfor %}
</div>
'''
    content = render_template_string(series_content, series=series, sections=sections)
    return render_template_string(BASE_TEMPLATE, title=f"Series {series.get('series_name')}", content=content)

@app.route('/raw/<series_id>')
def view_raw(series_id):
    try:
        raw_json = fetch_raw_json(series_id)
    except Exception as e:
        return f'Error: {e}', 500
    return Response(json.dumps(raw_json, ensure_ascii=False, indent=2), mimetype='application/json')

@app.route('/download/full/<series_id>/<lang>')
def download_full(series_id, lang):
    if lang not in ('en', 'hi'):
        return 'Invalid language', 400
    if not Document:
        return 'python-docx not installed', 500
    try:
        raw_json = fetch_raw_json(series_id)
    except Exception as e:
        return f'Error: {e}', 500
    cleaned = process_data(raw_json, lang)
    if lang == 'hi':
        cleaned = [e for e in cleaned if e['question_text'] or e.get('question_images')]
    if not cleaned:
        return 'No questions found for this test in the selected language.', 404
    from urllib.parse import urlparse
    url = QUESTION_API_PATTERN.format(series_id=series_id)
    parsed = urlparse(url)
    base_url = f"{parsed.scheme}://{parsed.netloc}"
    bio = create_docx(cleaned, base_url=base_url)
    filename = f"series_{series_id}_full_{lang}.docx"
    return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/download/section/<series_id>/<sid>/<lang>')
def download_section(series_id, sid, lang):
    if lang not in ('en', 'hi'):
        return 'Invalid language', 400
    if not Document:
        return 'python-docx not installed', 500
    try:
        raw_json = fetch_raw_json(series_id)
    except Exception as e:
        return f'Error: {e}', 500
    cleaned = process_data(raw_json, lang)
    if lang == 'hi':
        cleaned = [e for e in cleaned if e['question_text'] or e.get('question_images')]
    cleaned = [e for e in cleaned if str(e.get('section_id')) == str(sid)]
    if not cleaned:
        return 'No questions found for this section', 404
    from urllib.parse import urlparse
    url = QUESTION_API_PATTERN.format(series_id=series_id)
    parsed = urlparse(url)
    base_url = f"{parsed.scheme}://{parsed.netloc}"
    bio = create_docx(cleaned, base_url=base_url)
    filename = f"series_{series_id}_section_{sid}_{lang}.docx"
    return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/view/section/<series_id>/<sid>')
def view_section(series_id, sid):
    try:
        raw_json = fetch_raw_json(series_id)
    except Exception as e:
        return f'Error: {e}', 500
    cleaned = process_data(raw_json, 'en')
    sections = group_by_section(cleaned)
    sec = sections.get(sid)
    if not sec:
        return 'Section not found', 404
    view_content = '''
<h1 class="mb-3">Section: {{ section.name }} (Series ID: {{ series_id }})</h1>
<p><a href="/series/{{ series_id }}">&laquo; Back to series</a></p>
<ul>
{% for entry in questions %}
  <li><strong>[Q]</strong> {{ entry.question_text }}</li>
{% endfor %}
</ul>
'''
    content = render_template_string(view_content, section=sec, series_id=series_id, questions=sec['questions'])
    return render_template_string(BASE_TEMPLATE, title=f"Section {sec['name']}", content=content)

if __name__ == '__main__':
    print("WARNING: For production, use a WSGI server (like gunicorn) and set its timeout to at least 180 seconds for large downloads.")
    app.run(host='0.0.0.0', port=5000)
